import matplotlib.pyplot as plt
from data_loader import donors, blood_stock
from docx import Document
from docx.shared import Inches, Pt
import os

from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import ImageFormatter

# Generate Charts
def generate_charts():
    # 1. Blood Group Chart
    blood_groups = donors["blood_group"].value_counts()
    plt.figure(figsize=(8,5))
    plt.bar(blood_groups.index, blood_groups.values, color='skyblue')
    plt.title("Blood Group Distribution")
    plt.xlabel("Blood Group")
    plt.ylabel("Number of Donors")
    plt.savefig("blood_group_chart.png", bbox_inches='tight')
    plt.close()

    # 2. Blood Stock Chart
    plt.figure(figsize=(7,7))
    plt.pie(blood_stock["units_available"], labels=blood_stock["blood_group"], autopct="%1.1f%%", startangle=140)
    plt.title("Blood Stock Distribution")
    plt.savefig("blood_stock_chart.png", bbox_inches='tight')
    plt.close()

    # 3. City-wise Donors
    city = donors["city"].value_counts()
    plt.figure(figsize=(8,5))
    plt.bar(city.index, city.values, color='lightgreen')
    plt.title("City-wise Donors")
    plt.xlabel("City")
    plt.ylabel("Donors")
    plt.xticks(rotation=30)
    plt.savefig("city_donor_chart.png", bbox_inches='tight')
    plt.close()

    # 4. Donor Age Distribution
    plt.figure(figsize=(8,5))
    plt.hist(donors["age"], bins=6, color='coral', edgecolor='black')
    plt.title("Donor Age Distribution")
    plt.xlabel("Age")
    plt.ylabel("Frequency")
    plt.savefig("age_distribution_chart.png", bbox_inches='tight')
    plt.close()

def generate_code_screenshot(filename, output_png):
    with open(filename, 'r') as f:
        code = f.read()
    
    # We use ImageFormatter to create a PNG screenshot of the code
    # Adding line_numbers=True and a nice style
    formatter = ImageFormatter(font_size=16, line_numbers=True, style='monokai')
    with open(output_png, 'wb') as f:
        f.write(highlight(code, PythonLexer(), formatter))

def main():
    print("Generating charts...")
    generate_charts()
    
    print("Generating code screenshots...")
    # Just snapshotting some of the most important files
    files_to_screenshot = ['main.py', 'analysis.py', 'charts.py', 'data_loader.py']
    for file in files_to_screenshot:
        out_name = f"{file.split('.')[0]}_code.png"
        generate_code_screenshot(file, out_name)

    print("Creating Word document...")
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Report: Blood Bank Analytics Dashboard', 0)
    title.alignment = 1 # center

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This project is a Blood Bank Analytics Dashboard developed as part of a Python Summer Training program. "
        "It serves as a demonstration of data manipulation, analysis, and visualization concepts in Python. "
        "The application is built with a modular structure and features a console-based interactive menu."
    )

    doc.add_heading('2. Technologies Used', level=1)
    doc.add_paragraph("- Python: Core programming language.")
    doc.add_paragraph("- Pandas: Data manipulation and aggregation.")
    doc.add_paragraph("- NumPy: Numerical operations.")
    doc.add_paragraph("- Matplotlib: Data visualization (bar charts, pie charts, histograms).")

    doc.add_heading('3. Code Implementations (Screenshots)', level=1)
    doc.add_paragraph("Below are the screenshots of the code modules that power the application.")
    
    for file in files_to_screenshot:
        doc.add_heading(f"Module: {file}", level=2)
        out_name = f"{file.split('.')[0]}_code.png"
        if os.path.exists(out_name):
            doc.add_picture(out_name, width=Inches(6))

    doc.add_heading('4. Data Visualizations (Matplotlib Charts)', level=1)
    doc.add_paragraph("The following charts are generated using Matplotlib to visualize the dataset metrics.")

    doc.add_heading('Blood Group Distribution', level=2)
    doc.add_picture("blood_group_chart.png", width=Inches(5))
    
    doc.add_heading('Blood Stock Distribution', level=2)
    doc.add_picture("blood_stock_chart.png", width=Inches(5))
    
    doc.add_heading('City-wise Donors', level=2)
    doc.add_picture("city_donor_chart.png", width=Inches(5))
    
    doc.add_heading('Donor Age Distribution', level=2)
    doc.add_picture("age_distribution_chart.png", width=Inches(5))

    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "This project successfully applies Python programming concepts to solve a real-world data analytics problem. "
        "The modular architecture ensures the code is clean and readable, while Pandas and Matplotlib effectively "
        "handle data manipulation and visualization."
    )

    output_doc = "Project_Report.docx"
    doc.save(output_doc)
    print(f"Report saved to {output_doc}")

if __name__ == '__main__':
    main()
